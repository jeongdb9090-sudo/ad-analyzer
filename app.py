import base64
import io
import json
import os
import re
import time
import subprocess
from datetime import datetime
from zoneinfo import ZoneInfo
from urllib.parse import urlparse, parse_qs

# [추가] 히스토리 저장 시각이 서버 시간(대부분 UTC)이 아닌 한국시간(KST) 기준으로 기록되도록.
# 일부 서버 환경엔 IANA 시간대 데이터(tzdata)가 없을 수 있어, 실패 시 UTC+9 고정 오프셋으로 대체합니다.
try:
    KST = ZoneInfo("Asia/Seoul")
except Exception:
    from datetime import timezone, timedelta
    KST = timezone(timedelta(hours=9))
def now_kst():
    return datetime.now(KST)

# 필수 패키지 자동 설치 보장 함수 (맨 처음에 실행)
def ensure_package(package_name, import_name=None):
    if import_name is None:
        import_name = package_name
    try:
        __import__(import_name)
    except ImportError:
        try:
            subprocess.run(["pip", "install", package_name], check=True)
        except Exception:
            pass

ensure_package("requests")
ensure_package("Pillow", "PIL")
ensure_package("beautifulsoup4", "bs4")
ensure_package("google-generativeai", "google.generativeai")
ensure_package("openai")
ensure_package("anthropic")
ensure_package("playwright")
# --- [추가] 히스토리 영구 저장을 위한 구글시트 연동 패키지 (선택 사항, 미설정 시 자동 미사용) ---
ensure_package("tzdata")
ensure_package("gspread")
ensure_package("google-auth", "google.oauth2")
# --- [추가] 리포트를 워드/PDF 문서로 내보내기 위한 패키지 ---
ensure_package("python-docx", "docx")
ensure_package("fpdf2", "fpdf")

import requests
import streamlit as st
from PIL import Image
from bs4 import BeautifulSoup

# 스트림릿 클라우드 서버 환경에서 Playwright 브라우저 자동 설치 보장
@st.cache_resource
def install_playwright():
    try:
        subprocess.run(["playwright", "install", "chromium"], check=True)
    except Exception:
        pass

install_playwright()

try:
    from playwright.sync_api import sync_playwright
except Exception:
    pass

# ------------------------------------------------------------------
# 기본 설정 및 디자인 CSS
# ------------------------------------------------------------------
try:
    st.set_page_config(page_title="경쟁사 광고 소재 분석", layout="wide", page_icon="◆")
except Exception:
    pass

st.markdown("""
<style>
@import url('https://fonts.googleapis.com/css2?family=Space+Grotesk:wght@500;600;700&family=Inter:wght@400;500;600&family=JetBrains+Mono:wght@500&display=swap');

:root {
    --ink: #191B29;
    --muted: #555866;
    --paper: #FAF9F5;
    --surface: #FFFFFF;
    --border: #E2E1D9;
    --primary: #2A2F6B;
    --primary-soft: #EEEFF7;
    --amber: #D97706;
    --teal: #0D9488;
}

html, body, [class*="css"], .stMarkdown, p, label {
    font-family: 'Inter', sans-serif;
    color: var(--ink) !important;
}
/* [수정] span, div는 전역 강제 대상에서 제외했습니다. 태그 칩(멀티셀렉트 선택값)처럼
   색깔 있는 배경 위에 흰 글자가 필요한 요소들이 계속 이 규칙에 깔려서 안 보였기 때문입니다.
   대신 폰트만 통일 적용하고, 색상은 각 요소 상황에 맞게 아래에서 개별 지정합니다. */
html, body, [class*="css"], .stMarkdown, span, div {
    font-family: 'Inter', sans-serif;
}

.stApp { background-color: #FAF9F5 !important; }

h1, h2, h3, h4, h5, h6 {
    font-family: 'Space Grotesk', sans-serif !important;
    letter-spacing: -0.01em;
    color: var(--ink) !important;
}

.appbar { display: flex; align-items: center; gap: 14px; padding: 18px 4px 20px 4px; border-bottom: 1px solid var(--border); margin-bottom: 18px; }
.appbar-mark { width: 34px; height: 34px; border-radius: 8px; background: linear-gradient(135deg, var(--primary), var(--teal)); flex-shrink: 0; }
.appbar-title { font-family: 'Space Grotesk', sans-serif; font-weight: 700; font-size: 21px; line-height: 1.1; color: var(--ink) !important; }

.eyebrow { font-family: 'JetBrains Mono', monospace; font-size: 11px; letter-spacing: 0.08em; color: var(--primary) !important; background: var(--primary-soft); display: inline-block; padding: 3px 9px; border-radius: 4px; margin-bottom: 8px; font-weight: 600; }
.section-title { font-size: 20px; font-weight: 700; margin: 0 0 4px 0; color: var(--ink) !important; }
.section-desc { font-size: 13.5px; color: var(--muted) !important; margin-bottom: 18px; }

[data-testid="stSidebar"] {
    background-color: #F3F3EE !important;
    border-right: 1px solid var(--border) !important;
}
[data-testid="stSidebar"] * { color: var(--ink) !important; }

.stButton > button { border-radius: 8px; font-weight: 600; border: 1px solid var(--border); background-color: #FFFFFF !important; color: var(--ink) !important; }
.stButton > button[kind="primary"] { background-color: var(--primary) !important; color: #FFFFFF !important; border: none; }
.stButton > button[kind="primary"] * { color: #FFFFFF !important; }
.stButton > button[kind="primary"]:hover { background-color: #21245A !important; }

.score-grid { display: grid; grid-template-columns: repeat(3, 1fr); gap: 12px; margin: 12px 0; }
.score-card { border: 1px solid var(--border); border-radius: 8px; padding: 12px 14px; background-color: #FFFFFF !important; }
.score-cat { font-family: 'JetBrains Mono', monospace; font-size: 11px; letter-spacing: 0.05em; color: var(--muted) !important; text-transform: uppercase; font-weight: 600; }
.score-stars { color: var(--amber) !important; font-size: 16px; margin: 4px 0; }
.score-desc { font-size: 13px; color: var(--ink) !important; line-height: 1.5; white-space: pre-line; }

.comp-card { border: 1px solid var(--border); border-radius: 10px; padding: 14px 16px; background-color: #FFFFFF !important; margin-bottom: 10px; }
.comp-name { font-family: 'Space Grotesk', sans-serif; font-weight: 700; font-size: 15px; color: var(--ink) !important; }
.comp-meta { font-family: 'JetBrains Mono', monospace; font-size: 11px; color: var(--muted) !important; margin-top: 2px; }

div[data-baseweb="select"] > div,
div[data-baseweb="base-input"] > input,
[data-testid="stTextInput"] input,
[data-testid="stTextArea"] textarea,
[data-testid="stNumberInput"] input,
div[data-baseweb="popover"],
div[data-testid="stPopoverBody"],
div[data-baseweb="menu"],
div[data-baseweb="tag"],
div[data-baseweb="select"] {
    background-color: #FFFFFF !important;
    background: #FFFFFF !important;
    color: var(--ink) !important;
    border-color: var(--border) !important;
}

div[data-baseweb="select"] span,
div[data-baseweb="select"] div,
[data-baseweb="select"] * {
    background-color: transparent !important;
    color: var(--ink) !important;
}

div[data-baseweb="popover"] ul,
div[data-baseweb="popover"] li,
div[role="listbox"],
li[role="option"] {
    background-color: #FFFFFF !important;
    color: var(--ink) !important;
}
li[role="option"]:hover,
li[aria-selected="true"] {
    background-color: var(--primary-soft) !important;
    color: var(--ink) !important;
}

.align-bottom-btn { margin-top: 28px; }

/* ------------------------------------------------------------------
   [추가 CSS] 멀티셀렉트 선택 태그(칩) 글자를 흰색으로. 이제 위쪽의 전역 span 강제
   규칙을 제거했기 때문에, 이 규칙이 더 이상 다른 규칙에 덮이지 않고 정상 적용됩니다.
------------------------------------------------------------------ */
html body [data-testid="stMultiSelect"] span {
    color: #FFFFFF !important;
}

/* ------------------------------------------------------------------
   [추가 CSS - 2번] 다크모드에서 새는 요소 추가 보강
   config.toml에서 라이트 테마를 강제하지만, 팝업/파일업로더 등
   일부 컴포넌트는 브라우저 렌더링 타이밍에 따라 새는 경우가 있어 이중 방어.
------------------------------------------------------------------ */
[data-testid="stFileUploaderDropzone"],
[data-testid="stFileUploader"] section,
div[data-baseweb="popover"] * ,
ul[role="listbox"] li,
div[data-baseweb="calendar"],
div[data-baseweb="datepicker"] {
    background-color: #FFFFFF !important;
    color: var(--ink) !important;
}
input, textarea { caret-color: var(--ink) !important; }

/* ------------------------------------------------------------------
   [추가 CSS - 4번] 세그먼트 / 좌측 메뉴를 알약형·현대적 네비게이션으로
------------------------------------------------------------------ */
/* 세그먼트 선택 - 가로 알약(pill) 탭 */
div[data-testid="stSidebar"] div[data-testid="stRadio"]:has(div[role="radiogroup"]) {
    margin-bottom: 4px;
}
div[data-testid="stSidebar"] div[role="radiogroup"] {
    gap: 6px;
}
div[data-testid="stSidebar"] div[role="radiogroup"] > label {
    border: 1px solid var(--border) !important;
    background: #FFFFFF !important;
    border-radius: 999px;
    padding: 7px 14px !important;
    margin: 0 !important;
    transition: all 0.15s ease;
}
div[data-testid="stSidebar"] div[role="radiogroup"] > label:has(input:checked) {
    background: var(--primary) !important;
    border-color: var(--primary) !important;
}
div[data-testid="stSidebar"] div[role="radiogroup"] > label:has(input:checked) p {
    color: #FFFFFF !important;
    font-weight: 700;
}
div[data-testid="stSidebar"] div[role="radiogroup"] > label > div:first-child {
    display: none !important; /* 기본 라디오 원형 아이콘 숨김 */
}

/* 좌측 05단계 메뉴 - 세로 카드형 내비게이션 */
.nav-block div[role="radiogroup"] {
    flex-direction: column;
    gap: 4px;
}
.nav-block div[role="radiogroup"] > label {
    width: 100%;
    border-radius: 8px;
    padding: 10px 12px !important;
    border: 1px solid transparent !important;
    background: transparent !important;
}
.nav-block div[role="radiogroup"] > label:has(input:checked) {
    background: var(--primary-soft) !important;
    border-color: var(--primary) !important;
}
.nav-block div[role="radiogroup"] > label:has(input:checked) p {
    color: var(--primary) !important;
    font-weight: 700;
}
.nav-block div[role="radiogroup"] > label > div:first-child {
    display: none !important;
}
.sidebar-caption {
    font-family: 'JetBrains Mono', monospace;
    font-size: 10.5px;
    letter-spacing: 0.08em;
    color: var(--muted) !important;
    margin: 4px 0 6px 2px;
    text-transform: uppercase;
    font-weight: 700;
}
</style>
""", unsafe_allow_html=True)


def section_header(step, title, desc=""):
    st.markdown(f'<div class="eyebrow">STEP {step}</div>', unsafe_allow_html=True)
    st.markdown(f'<div class="section-title">{title}</div>', unsafe_allow_html=True)
    if desc:
        st.markdown(f'<div class="section-desc">{desc}</div>', unsafe_allow_html=True)


def stars(score, max_score=5):
    try:
        n = max(0, min(max_score, round(float(score))))
    except (ValueError, TypeError):
        n = 0
    return "★" * n + "☆" * (max_score - n)


# ------------------------------------------------------------------
# 세그먼트 & 메타 URL 사전
# ------------------------------------------------------------------
SEGMENTS = ["유아", "초등", "중등"]
DEFAULT_COMPETITORS = {
    "유아": ["윙크", "웅진스마트올", "밀크T아이", "리틀홈런"],
    "초등": ["밀크T", "아이스크림 홈런", "비상 온리원", "단꿈e", "기타"],
    "중등": ["밀크T중등", "웅진스마트올 중학", "비상 온리원 중등", "아이스크림 홈런 중등", "EBS"],
}

META_URL_MAP = {
    "윙크": "https://www.facebook.com/ads/library/?active_status=active&ad_type=all&country=KR&is_targeted_country=false&media_type=all&q=%EC%9C%99%ED%81%AC%ED%95%99%EC%8A%B5&search_type=keyword_unordered&sort_data[direction]=desc&sort_data[mode]=total_impressions",
    "웅진스마트올": "https://www.facebook.com/ads/library/?active_status=active&ad_type=all&country=KR&is_targeted_country=false&media_type=all&search_type=page&sort_data[direction]=desc&sort_data[mode]=total_impressions&view_all_page_id=100188454740792",
    "밀크T아이": "https://www.facebook.com/ads/library/?active_status=active&ad_type=all&country=KR&is_targeted_country=false&media_type=all&search_type=page&sort_data[direction]=desc&sort_data[mode]=total_impressions&view_all_page_id=113781314278271",
    "리틀홈런": "https://www.facebook.com/ads/library/?active_status=active&ad_type=all&country=KR&is_targeted_country=false&media_type=all&q=%EB%A6%AC%ED%8B%80%ED%99%88%EB%9F%B0&search_type=keyword_unordered&sort_data[direction]=desc&sort_data[mode]=total_impressions",
    "밀크T": "https://www.facebook.com/ads/library/?active_status=active&ad_type=all&country=KR&is_targeted_country=false&media_type=all&search_type=page&sort_data[direction]=desc&sort_data[mode]=total_impressions&view_all_page_id=699675066795625",
    "아이스크림 홈런": "https://www.facebook.com/ads/library/?active_status=active&ad_type=all&country=KR&is_targeted_country=false&media_type=all&search_type=page&sort_data[direction]=desc&sort_data[mode]=total_impressions&view_all_page_id=562669550571671",
    "비상 온리원": "https://www.facebook.com/ads/library/?active_status=active&ad_type=all&country=KR&is_targeted_country=false&media_type=all&search_type=page&sort_data[direction]=desc&sort_data[mode]=total_impressions&view_all_page_id=552773944780211",
    "단꿈e": "https://www.facebook.com/ads/library/?active_status=active&ad_type=all&country=KR&is_targeted_country=false&media_type=all&search_type=page&sort_data[direction]=desc&sort_data[mode]=total_impressions&view_all_page_id=350531981486027",
    "밀크T중등": "https://www.facebook.com/ads/library/?active_status=active&ad_type=all&country=KR&is_targeted_country=false&media_type=all&search_type=page&sort_data[direction]=desc&sort_data[mode]=total_impressions&view_all_page_id=101376489315136",
    "웅진스마트올 중학": "https://www.facebook.com/ads/library/?active_status=active&ad_type=all&country=KR&is_targeted_country=false&media_type=all&search_type=page&sort_data[direction]=desc&sort_data[mode]=total_impressions&view_all_page_id=103396781600446",
    "비상 온리원 중등": "https://www.facebook.com/ads/library/?active_status=active&ad_type=all&country=KR&is_targeted_country=false&media_type=all&search_type=page&sort_data[direction]=desc&sort_data[mode]=total_impressions&view_all_page_id=989750591106584",
    "아이스크림 홈런 중등": "https://www.facebook.com/ads/library/?active_status=active&ad_type=all&country=KR&is_targeted_country=false&media_type=all&q=%ED%99%88%EB%9F%B0%20%EC%A4%91%EB%93%B1&search_type=keyword_unordered&sort_data[direction]=desc&sort_data[mode]=total_impressions"
}

OWN_META_URL_MAP = {
    "유아": "https://www.facebook.com/ads/library/?active_status=active&ad_type=all&country=KR&is_targeted_country=false&media_type=all&search_type=page&sort_data[direction]=desc&sort_data[mode]=total_impressions&view_all_page_id=104085702734737",
    "초등": "https://www.facebook.com/ads/library/?active_status=active&ad_type=all&country=KR&is_targeted_country=false&media_type=all&search_type=page&sort_data[direction]=desc&sort_data[mode]=total_impressions&view_all_page_id=113924893334247",
    "중등": "https://www.facebook.com/ads/library/?active_status=active&ad_type=all&country=KR&is_targeted_country=false&media_type=all&search_type=page&sort_data[direction]=desc&sort_data[mode]=total_impressions&view_all_page_id=1600636653593633"
}

BASE_DIR = os.path.dirname(os.path.abspath(__file__))
HISTORY_FILE = os.path.join(BASE_DIR, "ad_signal_history.json")
BRAND_FILE = os.path.join(BASE_DIR, "ad_signal_brand.json")
COMPETITORS_FILE = os.path.join(BASE_DIR, "ad_signal_competitors.json")
PROFILES_FILE = os.path.join(BASE_DIR, "ad_signal_profiles.json")
SELECTORS_FILE = os.path.join(BASE_DIR, "ad_signal_selectors.json")
# --- [추가 - 5,6번] 자사 소재 분석 & 인사이트/갭분석을 세션이 아닌 파일(=구글시트)에 영구 저장 ---
OWN_FILE = os.path.join(BASE_DIR, "ad_signal_own.json")
WORK_STATE_FILE = os.path.join(BASE_DIR, "ad_signal_work_state.json")

DEFAULT_SELECTORS = {
    "ad_card_candidates": [
        'div[class*="_7jyg"]',
        'div[class*="_7j6g"]',
        'div[role="article"]',
        'div[data-testid="ad_library_card"]',
    ],
    "image_domain_keywords": ["scontent", "fbcdn"],
    "image_min_width": 150,
    "max_scroll_count": 15,
    "scroll_wait_ms": 2000,
    "initial_wait_timeout_ms": 15000,
}

def load_selectors():
    data = load_json(SELECTORS_FILE, None)
    if data is None:
        data = json.loads(json.dumps(DEFAULT_SELECTORS))
        save_json(SELECTORS_FILE, data)
        return data
    merged = dict(DEFAULT_SELECTORS)
    merged.update(data)
    return merged

def save_selectors(data): save_json(SELECTORS_FILE, data)
def reset_selectors():
    data = json.loads(json.dumps(DEFAULT_SELECTORS))
    save_json(SELECTORS_FILE, data)
    return data


# ------------------------------------------------------------------
# [추가 - 6번] 구글시트 기반 영구 저장소
# Streamlit Cloud의 로컬 파일시스템은 컨테이너가 재시작되면 초기화될 수 있어
# 히스토리가 계속 리셋되는 문제가 있었습니다. secrets.toml에
# GSHEET_ID / gcp_service_account 가 설정되어 있으면 자동으로 구글시트에
# 저장하고, 설정이 없으면 예전처럼 로컬 파일을 그대로 사용합니다(무설정 시 동작 100% 동일).
# 설정 방법은 채팅 답변 참고.
# ------------------------------------------------------------------
@st.cache_resource(show_spinner=False)
def _get_gsheet_client():
    try:
        import gspread
        from google.oauth2.service_account import Credentials
        if "gcp_service_account" not in st.secrets:
            return None
        scopes = ["https://www.googleapis.com/auth/spreadsheets"]
        creds = Credentials.from_service_account_info(dict(st.secrets["gcp_service_account"]), scopes=scopes)
        return gspread.authorize(creds)
    except Exception:
        return None

def _get_worksheet(sheet_name):
    client = _get_gsheet_client()
    if client is None:
        return None
    try:
        sheet_id = st.secrets.get("GSHEET_ID")
        if not sheet_id:
            return None
        sh = client.open_by_key(sheet_id)
        try:
            ws = sh.worksheet(sheet_name)
        except Exception:
            ws = sh.add_worksheet(title=sheet_name, rows=500, cols=1)
        return ws
    except Exception:
        return None

def load_json(path, default):
    """구글시트가 설정되어 있으면 시트에서, 아니면 로컬 파일에서 읽기"""
    sheet_name = os.path.splitext(os.path.basename(path))[0]
    ws = _get_worksheet(sheet_name)
    if ws is not None:
        try:
            col = ws.col_values(1)
            chunks = col[1:] if len(col) > 1 else []
            if chunks:
                return json.loads("".join(chunks))
            return default
        except Exception:
            return default
    if not os.path.exists(path): return default
    try:
        with open(path, "r", encoding="utf-8") as fp: return json.load(fp)
    except Exception: return default

def save_json(path, data):
    """구글시트가 설정되어 있으면 시트에, 아니면 로컬 파일에 쓰기
    (구글시트 셀당 50,000자 제한이 있어 4만자 단위로 청크 분할 저장)"""
    sheet_name = os.path.splitext(os.path.basename(path))[0]
    ws = _get_worksheet(sheet_name)
    if ws is not None:
        try:
            text = json.dumps(data, ensure_ascii=False)
            chunk_size = 40000
            chunks = [text[i:i + chunk_size] for i in range(0, len(text), chunk_size)] or [""]
            ws.clear()
            ws.update("A1", [["chunk"]] + [[c] for c in chunks])
            return
        except Exception:
            pass  # 실패 시 로컬 파일로 폴백
    with open(path, "w", encoding="utf-8") as fp: json.dump(data, fp, ensure_ascii=False, indent=2)


def load_history(): return load_json(HISTORY_FILE, [])
def save_history_entry(entry):
    history = load_history()
    history.insert(0, entry)
    save_json(HISTORY_FILE, history)

def load_all_brands(): return load_json(BRAND_FILE, {})
def save_brand(segment, data):
    all_brands = load_all_brands()
    all_brands[segment] = data
    save_json(BRAND_FILE, all_brands)

def load_competitors():
    data = load_json(COMPETITORS_FILE, {})
    changed = False
    for seg in SEGMENTS:
        if seg not in data:
            data[seg] = list(DEFAULT_COMPETITORS[seg])
            changed = True
    if changed: save_json(COMPETITORS_FILE, data)
    return data

def add_competitor(segment, name):
    data = load_competitors()
    if name and name not in data[segment]:
        data[segment].append(name)
        save_json(COMPETITORS_FILE, data)

def remove_competitor(segment, name):
    data = load_competitors()
    if name in data[segment]:
        data[segment].remove(name)
        save_json(COMPETITORS_FILE, data)

def load_all_profiles(): return load_json(PROFILES_FILE, {})
def save_profile_entry(segment, competitor, entry):
    data = load_all_profiles()
    data.setdefault(segment, {}).setdefault(competitor, []).insert(0, entry)
    save_json(PROFILES_FILE, data)

# --- [추가 - 5,6번] 자사 소재 분석 결과 영구 저장 (기존엔 세션에만 있어서 새로고침하면 사라짐) ---
def load_own_analysis(segment):
    return load_json(OWN_FILE, {}).get(segment)

def save_own_analysis(segment, data):
    all_own = load_json(OWN_FILE, {})
    all_own[segment] = data
    save_json(OWN_FILE, all_own)

# --- [추가 - 5,6번] 세그먼트별 작업 상태(인사이트/갭분석/스토리보드) 영구 저장 ---
def load_work_state(segment):
    default = {"insight": "", "gap_analysis": "", "ideas": ""}
    default.update(load_json(WORK_STATE_FILE, {}).get(segment, {}))
    return default

def save_work_state(segment, data):
    all_state = load_json(WORK_STATE_FILE, {})
    all_state[segment] = data
    save_json(WORK_STATE_FILE, all_state)


# ------------------------------------------------------------------
# Playwright 크롤링 함수
# ------------------------------------------------------------------
_EXTRACT_ADS_JS = """
(config) => {
    const { ad_card_candidates, image_domain_keywords, image_min_width } = config;
    let cards = [];
    let usedSelector = null;
    for (const sel of ad_card_candidates) {
        try {
            const found = document.querySelectorAll(sel);
            if (found.length > 0) {
                cards = Array.from(found);
                usedSelector = sel;
                break;
            }
        } catch (e) {}
    }
    const scope = cards.length ? cards : [document.body];
    if (!cards.length) usedSelector = 'FALLBACK: document.body';

    // [추가 - 3번] 영상 소재는 이미지가 아니라서 수집 대상에서 자동 제외되는데,
    // "왜 개수가 안 맞지?" 헷갈리지 않도록 몇 건이 영상이라 제외됐는지 별도로 셉니다.
    const seenVideoCards = new Set();
    scope.forEach((card, idx) => {
        const hasVideoTag = card.querySelector('video') !== null;
        const hasVideoAttr = card.querySelector('[aria-label*="video" i], [aria-label*="동영상"]') !== null;
        if (hasVideoTag || hasVideoAttr) seenVideoCards.add(idx);
    });
    const videoCount = seenVideoCards.size;

    const seen = new Set();
    const items = [];
    scope.forEach((card) => {
        const imgs = card.querySelectorAll('img');
        imgs.forEach((img) => {
            const src = img.currentSrc || img.src || '';
            if (!src || seen.has(src)) return;

            let renderedWidth = img.naturalWidth || 0;
            let renderedHeight = img.naturalHeight || 0;
            if (!renderedWidth) {
                try { 
                    const rect = img.getBoundingClientRect();
                    renderedWidth = Math.round(rect.width) || 0;
                    renderedHeight = Math.round(rect.height) || 0;
                } catch (e) {}
            }
            if (!renderedWidth) renderedWidth = img.width || 0;
            if (!renderedHeight) renderedHeight = img.height || 0;

            if (renderedWidth > 0 && renderedWidth < image_min_width) return;
            if (renderedHeight > 0 && renderedHeight < image_min_width) return;
            
            if (renderedWidth > 0 && renderedHeight > 0) {
                const ratio = renderedWidth / renderedHeight;
                if (ratio > 0.8 && ratio < 1.2 && renderedWidth < 180) return;
            }

            const matchesDomain = image_domain_keywords.some((k) => src.includes(k));
            if (!matchesDomain) return;

            seen.add(src);
            let bodyText = '';
            try { bodyText = (card.innerText || '').slice(0, 300); } catch (e) {}
            items.push({ src, bodyText });
        });
    });
    return { usedSelector, cardCount: cards.length, items, videoCount };
}
"""

def _upgrade_image_resolution(url):
    if not url: return url
    try:
        def _bump(match):
            w, h = int(match.group(1)), int(match.group(2))
            return f"s{min(int(w * 1.6), 1080)}x{min(int(h * 1.6), 1080)}"
        return re.sub(r"s(\d{2,4})x(\d{2,4})", _bump, url)
    except Exception:
        return url

def _scrape_once(library_url, max_items, cfg, status_callback=None):
    results = []
    debug_info = {"used_selector": None, "card_count": 0, "video_count": 0}
    
    if status_callback: status_callback("브라우저 엔진을 구동하고 있습니다...")
    with sync_playwright() as p:
        browser = p.chromium.launch(
            headless=True,
            args=["--disable-blink-features=AutomationControlled", "--no-sandbox", "--disable-setuid-sandbox", "--disable-dev-shm-usage", "--disable-gpu", "--disable-extensions", "--no-zygote"]
        )
        try:
            context = browser.new_context(
                user_agent="Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/124.0.0.0 Safari/537.36",
                viewport={"width": 1280, "height": 800},
                locale="ko-KR"
            )
            page = context.new_page()
            page.add_init_script("Object.defineProperty(navigator, 'webdriver', {get: () => undefined});")
            page.route("**/*", lambda route: route.abort() if route.request.resource_type in ("media", "font") else route.continue_())
            
            if status_callback: status_callback("메타 광고 라이브러리 페이지에 접속 중입니다...")
            page.goto(library_url, timeout=45000, wait_until="domcontentloaded")
            time.sleep(5)
            if page.is_closed(): raise RuntimeError("브라우저 종료됨")
            try:
                page.wait_for_selector(", ".join(cfg["ad_card_candidates"]), timeout=cfg.get("initial_wait_timeout_ms", 15000))
            except Exception:
                pass
            
            if status_callback: status_callback("활성 광고 데이터를 스크롤하여 불러오는 중...")
            last_height = 0
            for i in range(cfg.get("max_scroll_count", 8)):
                if page.is_closed(): break
                page.evaluate("window.scrollTo(0, document.body.scrollHeight)")
                time.sleep(cfg.get("scroll_wait_ms", 2000) / 1000)
                new_height = page.evaluate("document.body.scrollHeight")
                if new_height == last_height: break
                last_height = new_height
            
            time.sleep(1.5)
            if status_callback: status_callback("광고 소재 이미지 및 텍스트를 추출하고 있습니다...")
            extracted = page.evaluate(_EXTRACT_ADS_JS, cfg)
            debug_info["used_selector"] = extracted.get("usedSelector")
            debug_info["card_count"] = extracted.get("cardCount", 0)
            debug_info["video_count"] = extracted.get("videoCount", 0)  # [추가 - 3번]
            
            total_items = extracted.get("items", [])[:max_items]
            for idx, item in enumerate(total_items, start=1):
                if status_callback: status_callback(f"고화질 광고 소재 이미지 다운로드 중 ({idx}/{len(total_items)}건)...")
                img_url = item.get("src")
                body_text = item.get("bodyText", "")
                img_bytes = None
                if img_url:
                    hi_res = _upgrade_image_resolution(img_url)
                    for cand in [hi_res, img_url]:
                        try:
                            resp = requests.get(cand, headers={"User-Agent": "Mozilla/5.0"}, timeout=8)
                            if resp.status_code == 200:
                                img_bytes = resp.content
                                break
                        except Exception:
                            continue
                results.append({"id": f"pw_{idx}_{time.time()}", "fn": f"ad_{idx}.png", "bytes": img_bytes, "body": body_text, "snapshot_url": library_url})
        finally:
            try: browser.close()
            except Exception: pass
    return results, debug_info

def scrape_meta_ads_with_playwright(library_url, max_items=12, selectors=None, status_callback=None, _retry=True):
    cfg = selectors or load_selectors()
    try:
        results, debug_info = _scrape_once(library_url, max_items, cfg, status_callback)
        return results, debug_info, None
    except Exception as e:
        err_text = str(e)
        if ("closed" in err_text or "종료" in err_text) and _retry:
            time.sleep(2)
            return scrape_meta_ads_with_playwright(library_url, max_items, cfg, status_callback, _retry=False)
        return [], {"used_selector": None, "card_count": 0}, f"오류 발생: {err_text}"

def create_image_grid_collage(images_bytes_list, cols=4, thumb_size=(180, 180), max_images=12):
    try:
        pil_images = []
        for b in images_bytes_list[:max_images]:
            if not b: continue
            try:
                img = Image.open(io.BytesIO(b)).convert("RGB")
                img.thumbnail(thumb_size)
                pil_images.append(img)
            except Exception: pass
        if not pil_images: return None
        num_imgs = len(pil_images)
        rows = (num_imgs + cols - 1) // cols
        collage = Image.new("RGB", (cols * thumb_size[0], rows * thumb_size[1]), (255, 255, 255))
        for idx, img in enumerate(pil_images):
            collage.paste(img, ((idx % cols) * thumb_size[0], (idx // cols) * thumb_size[1]))
        buf = io.BytesIO()
        collage.save(buf, format="JPEG", quality=65)
        return buf.getvalue()
    except Exception: return None


def create_split_collage(top_images_bytes, bottom_images_bytes, cols=4, thumb_size=(160, 160),
                          top_max=16, bottom_max=12):
    """
    [추가 - 통합분석] 경쟁사(여러 브랜드 통합) 콜라주를 위쪽에, 자사 콜라주를 아래쪽에 놓고
    가운데 굵은 컬러 구분선을 넣어 하나의 이미지로 합칩니다.
    이렇게 하면 AI 호출 1번으로 두 그룹을 동시에 보여주고 비교시킬 수 있어 무료 API 한도를 아낄 수 있습니다.
    """
    try:
        def _make_block(images_bytes, max_n):
            pil_images = []
            for b in images_bytes[:max_n]:
                if not b: continue
                try:
                    img = Image.open(io.BytesIO(b)).convert("RGB")
                    img.thumbnail(thumb_size)
                    pil_images.append(img)
                except Exception: pass
            if not pil_images: return None
            rows = (len(pil_images) + cols - 1) // cols
            block = Image.new("RGB", (cols * thumb_size[0], rows * thumb_size[1]), (255, 255, 255))
            for idx, img in enumerate(pil_images):
                block.paste(img, ((idx % cols) * thumb_size[0], (idx // cols) * thumb_size[1]))
            return block

        top_block = _make_block(top_images_bytes, top_max)
        bottom_block = _make_block(bottom_images_bytes, bottom_max)
        if top_block is None and bottom_block is None: return None

        width = cols * thumb_size[0]
        divider_h = 26
        top_h = top_block.height if top_block else 0
        bottom_h = bottom_block.height if bottom_block else 0
        total_h = top_h + (divider_h if top_block and bottom_block else 0) + bottom_h
        combined = Image.new("RGB", (width, total_h), (255, 255, 255))
        y = 0
        if top_block:
            combined.paste(top_block, (0, 0))
            y += top_h
        if top_block and bottom_block:
            from PIL import ImageDraw
            draw = ImageDraw.Draw(combined)
            draw.rectangle([0, y, width, y + divider_h], fill=(217, 119, 6))  # 오렌지색 구분선(위=경쟁사, 아래=자사)
            y += divider_h
        if bottom_block:
            combined.paste(bottom_block, (0, y))

        buf = io.BytesIO()
        combined.save(buf, format="JPEG", quality=70)
        return buf.getvalue()
    except Exception:
        return None


COMBINED_GAP_ANALYSIS_PROMPT = """당신은 수석 브랜드 전략 컨설턴트 겸 퍼포먼스 마케팅 분석가입니다.
첨부한 이미지는 위/아래 두 구역으로 나뉘어 있고, 가운데 오렌지색 굵은 줄로 구분되어 있습니다.
- **이미지의 위쪽 구역**: 여러 경쟁사들({competitor_names})의 광고 소재를 한데 모은 콜라주입니다. 개별 브랜드 구분 없이, 경쟁사 그룹 전체가 공통적으로 어떤 메시지/비주얼 전략을 쓰고 있는지 통합적으로 파악해주세요.
- **이미지의 아래쪽 구역**: 자사({own_brand_label}) 광고 소재를 모은 콜라주입니다.

두 구역을 비교 분석해서 아래 형식으로 정리해주세요. 각 항목은 실제 이미지에서 관찰되는 내용에 근거해서 구체적으로 작성해주세요.

### 1. 경쟁사 공통 위닝 포인트 (3가지)
1. **[키워드]**: 설명...
2. **[키워드]**: 설명...
3. **[키워드]**: 설명...

### 2. 자사 소재 강점 / 아쉬운 점
- 👍 강점: ...
- 👎 아쉬운 점: ...

### 3. 경쟁사는 다루지만 자사 소재에는 부족한 메시지
- ...

### 4. 보강하면 좋을 메시지 (우선순위 순, 이유 포함)
1. ...
2. ...
3. ...

### 5. 비주얼/톤앤매너 측면에서 참고할 점
- ...

### 6. 우리만 갖고 있는 강점 (계속 유지할 것)
- ...
"""

COMPETITOR_ONLY_TREND_PROMPT = """당신은 수석 브랜드 전략가입니다. 첨부한 이미지는 여러 경쟁사({competitor_names})의
광고 소재를 한데 모은 콜라주입니다. 개별 브랜드 구분 없이 경쟁사 그룹 전체가 공통적으로 활용하는
성공 패턴(위닝 포인트)을 아래 형식으로 정리해주세요.

### 핵심 위닝 포인트 3가지
1. **[키워드]**: 설명...
2. **[키워드]**: 설명...
3. **[키워드]**: 설명...

### 종합 마케팅 인사이트
- 시장 내 공통적인 소구 트렌드 및 시사점
"""

_SCORE_LABEL_TO_KEY = {
    "메시지_좋은점": "msg_good", "메시지_아쉬운점": "msg_bad", "메시지_평점": "msg_score",
    "비주얼_좋은점": "vis_good", "비주얼_아쉬운점": "vis_bad", "비주얼_평점": "vis_score",
}
_SCORE_EMPTY = {k: "" for k in _SCORE_LABEL_TO_KEY.values()}

def parse_labeled_text(text, label_to_key, empty_fields):
    fields = dict(empty_fields)
    current = None
    for raw_line in text.splitlines():
        line = raw_line.strip()
        if not line: continue
        matched_key = None
        for label, key in label_to_key.items():
            if line.startswith(label + ":") or line.startswith(label + " :"):
                fields[key] = line.split(":", 1)[1].strip() if ":" in line else ""
                current = key
                matched_key = key
                break
        if matched_key is None and current:
            fields[current] = (fields[current] + " " + line).strip()
    return fields

def parse_integrated_report(text): return parse_labeled_text(text, _SCORE_LABEL_TO_KEY, _SCORE_EMPTY)

def render_integrated_scorecard(report):
    msg_desc = f"👍 **장점**: {report.get('msg_good', '')}\n👎 **아쉬운점**: {report.get('msg_bad', '')}"
    vis_desc = f"👍 **장점**: {report.get('vis_good', '')}\n👎 **아쉬운점**: {report.get('vis_bad', '')}"
    overall_desc = report.get('overall_desc', '')
    cats = [
        ("메시지 전략", msg_desc, report.get("msg_score", 0)),
        ("비주얼 / 디자인", vis_desc, report.get("vis_score", 0)),
        ("종합 타겟 전달력 평가", overall_desc, report.get("overall_score", 0)),
    ]
    html = '<div class="score-grid">'
    for label, desc, score in cats:
        html += f'<div class="score-card"><div class="score-cat">{label}</div><div class="score-stars">{stars(score)}</div><div class="score-desc">{desc}</div></div>'
    html += "</div>"
    st.markdown(html, unsafe_allow_html=True)

class QuotaExceededError(Exception):
    """무료 티어 쿼터/레이트리밋(429) 초과 - 재시도로도 해결 안 될 때 사용자에게 안내할 전용 예외"""
    pass


def _extract_retry_seconds(err_text, default=15):
    """구글 API 에러 메시지에 담긴 'retry_delay { seconds: N }' 값을 파싱, 없으면 기본값 사용"""
    m = re.search(r"retry_delay\s*\{\s*seconds:\s*(\d+)", err_text)
    if m: return int(m.group(1)) + 1
    m2 = re.search(r"retry in ([\d.]+)s", err_text, re.IGNORECASE)
    if m2: return int(float(m2.group(1))) + 1
    return default


def _call_gemini(api_key, prompt_text, collage_bytes):
    try:
        import google.generativeai as genai
    except ImportError:
        subprocess.run(["pip", "install", "google-generativeai"], check=True)
        import google.generativeai as genai
    genai.configure(api_key=api_key)
    # [수정 - 3번] gemini-1.5-pro는 완전히 서비스 종료(404)됨.
    # -latest 별칭을 사용하면 구글이 모델을 교체해도 자동으로 최신 모델을 가리켜서
    # 앞으로 이런 단종 문제가 재발할 확률이 낮습니다.
    model = genai.GenerativeModel("gemini-flash-latest")
    contents = [prompt_text]
    if collage_bytes: contents.append(Image.open(io.BytesIO(collage_bytes)))
    return model.generate_content(contents).text


def _call_openai(api_key, prompt_text, collage_bytes):
    from openai import OpenAI
    client = OpenAI(api_key=api_key)
    content_payload = [{"type": "text", "text": prompt_text}]
    if collage_bytes:
        content_payload.append({"type": "image_url", "image_url": {"url": f"data:image/jpeg;base64,{base64.b64encode(collage_bytes).decode('utf-8')}"}})
    # [수정 - 3번] gpt-4o -> gpt-5.1 (최신 플래그십, 비전 지원)
    return client.chat.completions.create(model="gpt-5.1", messages=[{"role": "user", "content": content_payload}]).choices[0].message.content


def _call_claude(api_key, prompt_text, collage_bytes):
    import anthropic
    client = anthropic.Anthropic(api_key=api_key)
    content_payload = []
    if collage_bytes:
        content_payload.append({"type": "image", "source": {"type": "base64", "media_type": "image/jpeg", "data": base64.b64encode(collage_bytes).decode('utf-8')}})
    content_payload.append({"type": "text", "text": prompt_text})
    # [수정 - 3번] claude-3-5-sonnet-20241022 -> claude-sonnet-5 (최신 모델)
    return client.messages.create(model="claude-sonnet-5", max_tokens=2000, messages=[{"role": "user", "content": content_payload}]).content[0].text


def run_unified_ai_prompt(ai_provider, api_key, prompt_text, collage_bytes=None, status_callback=None, max_retries=3):
    """
    [수정 - 2번] 무료 API 요금제는 분당 요청 횟수가 매우 낮게 제한되어 있어(예: Gemini 무료 티어
    분당 5회) '429 quota exceeded' 오류가 자주 발생합니다. 이 함수는 그 오류를 감지하면
    에러 메시지 안의 재시도 대기 시간(retry_delay)만큼 자동으로 기다렸다가 최대 max_retries번
    다시 시도합니다. 진행 상황은 status_callback으로 화면에 실시간 보고합니다.
    """
    dispatch = {
        "Gemini (Google)": _call_gemini,
        "ChatGPT (OpenAI)": _call_openai,
        "Claude (Anthropic)": _call_claude,
    }
    fn = dispatch.get(ai_provider)
    if fn is None:
        raise ValueError(f"알 수 없는 AI 엔진: {ai_provider}")

    last_err_text = ""
    for attempt in range(1, max_retries + 1):
        if status_callback:
            status_callback(f"{ai_provider} 모델을 호출하는 중입니다... (시도 {attempt}/{max_retries})")
        try:
            result_text = fn(api_key, prompt_text, collage_bytes)
            if status_callback: status_callback("응답을 받았습니다. 리포트로 정리하는 중...")
            return result_text
        except Exception as e:
            err_text = str(e)
            last_err_text = err_text
            is_quota_error = ("429" in err_text or "quota" in err_text.lower() or "rate" in err_text.lower())
            if is_quota_error and attempt < max_retries:
                wait_s = _extract_retry_seconds(err_text)
                if status_callback:
                    status_callback(f"⏳ 무료 API 사용량 한도에 걸렸습니다. {wait_s}초 후 자동 재시도합니다... ({attempt}/{max_retries})")
                time.sleep(wait_s)
                continue
            elif is_quota_error:
                raise QuotaExceededError(
                    f"{ai_provider}의 무료 API 사용량 한도를 초과했습니다. "
                    f"잠시 후 다시 시도하시거나, Google AI Studio에서 결제(유료 티어)를 활성화하면 "
                    f"한도가 크게 늘어납니다. (원본 오류: {err_text[:200]})"
                )
            else:
                raise
    raise QuotaExceededError(f"{max_retries}회 재시도했지만 계속 실패했습니다: {last_err_text[:200]}")


def run_brand_integrated_analysis(ai_provider, api_key, brand_name, images_bytes_list, status_callback=None):
    """(현재는 03번 탭에서 사용하지 않는 레거시 함수 - 필요 시 개별 브랜드 분석용으로 재사용 가능)"""
    if status_callback: status_callback("소재 이미지를 하나의 콜라주로 합치는 중...")
    collage = create_image_grid_collage(images_bytes_list)
    text = run_unified_ai_prompt(
        ai_provider, api_key,
        BRAND_INTEGRATED_ANALYSIS_PROMPT.format(brand_name=brand_name),
        collage, status_callback=status_callback,
    )
    if status_callback: status_callback("리포트 항목을 파싱하는 중...")
    return parse_integrated_report(text)


def run_combined_gap_analysis(ai_provider, api_key, competitor_images, own_images, competitor_names, own_brand_label, status_callback=None):
    """
    [핵심 - 03번 탭] 경쟁사 여러 브랜드의 소재를 하나로, 자사 소재를 하나로 인식시켜
    AI 호출 단 1번으로 '경쟁사 공통 위닝포인트 + 자사 강약점 + 메시지 갭'을 모두 받아옵니다.
    """
    if status_callback: status_callback("경쟁사 소재 콜라주 + 자사 소재 콜라주를 하나의 이미지로 합치는 중...")
    combined_img = create_split_collage(competitor_images, own_images)
    names_str = ", ".join(competitor_names) if competitor_names else "경쟁사"
    prompt = COMBINED_GAP_ANALYSIS_PROMPT.format(competitor_names=names_str, own_brand_label=own_brand_label)
    return run_unified_ai_prompt(ai_provider, api_key, prompt, combined_img, status_callback=status_callback)


def run_competitor_trend_only(ai_provider, api_key, competitor_images, competitor_names, status_callback=None):
    """자사 소재가 아직 없을 때, 경쟁사 소재만으로 트렌드 인사이트를 뽑는 보조 함수 (AI 호출 1번)"""
    if status_callback: status_callback("경쟁사 소재를 하나의 콜라주로 합치는 중...")
    collage = create_image_grid_collage(competitor_images, max_images=16)
    names_str = ", ".join(competitor_names) if competitor_names else "경쟁사"
    prompt = COMPETITOR_ONLY_TREND_PROMPT.format(competitor_names=names_str)
    return run_unified_ai_prompt(ai_provider, api_key, prompt, collage, status_callback=status_callback)


def gather_collected_materials(segment):
    """01/02번 탭에서 세션에 모아둔 이미지를 브랜드별로 취합 (03/04번 탭 공용)"""
    competitors = load_competitors().get(segment, [])
    comp_materials = {}
    for c in competitors:
        items = st.session_state.get(f"comp_items_{c}", [])
        byte_list = [it["bytes"] for it in items if it.get("bytes")]
        if byte_list:
            comp_materials[c] = byte_list
    own_key = f"own_items_자사({segment})"
    own_items = st.session_state.get(own_key, [])
    own_bytes = [it["bytes"] for it in own_items if it.get("bytes")]
    return comp_materials, own_bytes


# ------------------------------------------------------------------
# [추가] 리포트 문서 내보내기 (DOCX / PDF)
# PDF는 한글 폰트가 없으면 글자가 깨지므로, 최초 1회 Noto Sans KR 폰트를
# 구글 폰트 공식 저장소(raw.githubusercontent.com)에서 내려받아 캐시해 사용합니다.
# ------------------------------------------------------------------
_KOREAN_FONT_URL = "https://raw.githubusercontent.com/google/fonts/main/ofl/notosanskr/NotoSansKR%5Bwght%5D.ttf"
_KOREAN_FONT_PATH = os.path.join(BASE_DIR, "NotoSansKR.ttf")

@st.cache_resource(show_spinner=False)
def _ensure_korean_font():
    if os.path.exists(_KOREAN_FONT_PATH):
        return _KOREAN_FONT_PATH
    try:
        resp = requests.get(_KOREAN_FONT_URL, timeout=20)
        if resp.status_code == 200:
            with open(_KOREAN_FONT_PATH, "wb") as f:
                f.write(resp.content)
            return _KOREAN_FONT_PATH
    except Exception:
        pass
    return None


def _add_markdown_to_docx(doc, text):
    """간단한 마크다운(#, ##, ###, **볼드**)을 워드 문단/제목 스타일로 변환"""
    for raw_line in (text or "").split("\n"):
        line = raw_line.rstrip()
        if not line.strip():
            doc.add_paragraph("")
            continue
        if line.startswith("### "):
            doc.add_heading(line[4:].strip(), level=3)
        elif line.startswith("## "):
            doc.add_heading(line[3:].strip(), level=2)
        elif line.startswith("# "):
            doc.add_heading(line[2:].strip(), level=1)
        else:
            p = doc.add_paragraph()
            parts = re.split(r"(\*\*.+?\*\*)", line)
            for part in parts:
                if part.startswith("**") and part.endswith("**"):
                    run = p.add_run(part[2:-2])
                    run.bold = True
                else:
                    p.add_run(part)


def build_report_docx(segment, brand_info, insight_text, gap_text, ideas_text):
    from docx import Document
    doc = Document()
    doc.add_heading(f"{segment} 경쟁사 광고 소재 분석 리포트", level=1)
    doc.add_paragraph(f"생성일시: {now_kst().strftime('%Y-%m-%d %H:%M')}")

    if brand_info and any(brand_info.values()):
        doc.add_heading("브랜드 정보", level=2)
        label_map = {"brand_name": "브랜드/제품명", "brand_product": "제품 설명", "brand_usp": "핵심 USP",
                     "target_audience": "타겟 고객", "brand_design_memory": "디자인 메모리"}
        for k, v in brand_info.items():
            if v: doc.add_paragraph(f"{label_map.get(k, k)}: {v}")

    if insight_text:
        doc.add_heading("경쟁사 트렌드 인사이트", level=2)
        _add_markdown_to_docx(doc, insight_text)
    if gap_text:
        doc.add_heading("메시지 갭 분석", level=2)
        _add_markdown_to_docx(doc, gap_text)
    if ideas_text:
        doc.add_heading("신규 소재 스토리보드 아이디어", level=2)
        _add_markdown_to_docx(doc, ideas_text)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def build_report_pdf(segment, brand_info, insight_text, gap_text, ideas_text):
    from fpdf import FPDF
    font_path = _ensure_korean_font()
    pdf = FPDF()
    pdf.set_auto_page_break(auto=True, margin=15)
    pdf.add_page()
    if font_path:
        pdf.add_font("Korean", "", font_path)
        pdf.set_font("Korean", size=16)
    else:
        pdf.set_font("Helvetica", size=16)  # 폰트 다운로드 실패 시 폴백 (한글이 깨질 수 있음)

    pdf.multi_cell(0, 10, f"{segment} 경쟁사 광고 소재 분석 리포트")
    pdf.set_font_size(10)
    pdf.multi_cell(0, 6, f"생성일시: {now_kst().strftime('%Y-%m-%d %H:%M')}")
    pdf.ln(4)

    def _section(title, body):
        if not body: return
        pdf.set_font_size(13)
        pdf.multi_cell(0, 8, title)
        pdf.set_font_size(10)
        pdf.multi_cell(0, 6, body.replace("**", "").replace("### ", "").replace("## ", "").replace("# ", ""))
        pdf.ln(3)

    if brand_info and any(brand_info.values()):
        label_map = {"brand_name": "브랜드/제품명", "brand_product": "제품 설명", "brand_usp": "핵심 USP",
                     "target_audience": "타겟 고객", "brand_design_memory": "디자인 메모리"}
        info_lines = "\n".join([f"{label_map.get(k, k)}: {v}" for k, v in brand_info.items() if v])
        _section("■ 브랜드 정보", info_lines)
    _section("■ 경쟁사 트렌드 인사이트", insight_text)
    _section("■ 메시지 갭 분석", gap_text)
    _section("■ 신규 소재 스토리보드 아이디어", ideas_text)

    return bytes(pdf.output())


def render_export_buttons(segment, brand_info, insight_text, gap_text, ideas_text, key_prefix):
    """DOCX/PDF 다운로드 버튼 2개를 나란히 표시"""
    exp_cols = st.columns(2)
    with exp_cols[0]:
        try:
            docx_bytes = build_report_docx(segment, brand_info, insight_text, gap_text, ideas_text)
            st.download_button(
                "📄 워드(.docx)로 저장", data=docx_bytes,
                file_name=f"{segment}_ad_report.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                key=f"{key_prefix}_docx_dl",
            )
        except Exception as e:
            st.caption(f"워드 생성 실패: {e}")
    with exp_cols[1]:
        try:
            pdf_bytes = build_report_pdf(segment, brand_info, insight_text, gap_text, ideas_text)
            st.download_button(
                "📕 PDF로 저장", data=pdf_bytes,
                file_name=f"{segment}_ad_report.pdf",
                mime="application/pdf",
                key=f"{key_prefix}_pdf_dl",
            )
        except Exception as e:
            st.caption(f"PDF 생성 실패: {e}")


# ------------------------------------------------------------------
# [소재 수집 UI] - [수정] 이 단계에서는 AI 분석을 하지 않고 '수집'만 합니다.
# 실제 AI 분석(경쟁사 통합 + 자사 통합 + 갭 비교)은 03번 탭에서 단 1회의 API 호출로 처리해서
# 무료 API 요금제의 분당 요청 한도를 아끼도록 구조를 바꿨습니다.
# ------------------------------------------------------------------
def render_material_section(prefix, selected_comp, default_url, max_items=30):
    tab1, tab2 = st.tabs(["🚀 자동 크롤링 수집", "📁 예비 업로드"])
    sess_key = f"{prefix}_items_{selected_comp}"
    if sess_key not in st.session_state: st.session_state[sess_key] = []

    with tab1:
        meta_url = st.text_input("메타 광고 라이브러리 페이지 URL", value=default_url, key=f"{prefix}_meta_url_input_{selected_comp}")
        if st.button("🚀 전체 라이브 소재 자동 수집 실행", key=f"{prefix}_crawl_btn_{selected_comp}", type="primary"):
            if not meta_url.strip():
                st.warning("메타 라이브러리 URL을 입력해주세요.")
            else:
                with st.status(f"'{selected_comp}' 광고 라이브러리 수집 진행 중...", expanded=True) as status_box:
                    def update_status(msg):
                        status_box.update(label=msg, state="running")

                    ads, debug_info, err = scrape_meta_ads_with_playwright(
                        meta_url.strip(), max_items=max_items, selectors=load_selectors(), status_callback=update_status
                    )

                    if err:
                        status_box.update(label="수집 중 오류 발생", state="error")
                        st.error(err)
                    elif not ads:
                        status_box.update(label="수집된 활성 광고가 없습니다.", state="complete")
                        st.info(f"⚠️ [{selected_comp}] 활성 광고를 수집하지 못했습니다. [예비 업로드] 탭을 활용해 주세요.")
                        st.session_state[sess_key] = []
                    else:
                        # [수정 - 3번] 전체 카드 수 대비 이미지로 수집된 건수, 영상이라 제외된 건수를 함께 안내
                        video_count = debug_info.get("video_count", 0)
                        card_count = debug_info.get("card_count", 0)
                        status_box.update(label=f"'{selected_comp}' 총 {len(ads)}건 수집 완료!", state="complete")
                        st.session_state[sess_key] = ads
                        msg = f"[{selected_comp}] 이미지 소재 {len(ads)}건 수집 완료"
                        if video_count > 0:
                            msg += f" · 동영상 소재 {video_count}건은 이미지가 아니라서 자동 제외됨"
                        st.success(msg)
                        st.caption(f"페이지에서 감지된 전체 광고 카드 {card_count}개 중 이미지 소재만 추출했습니다.")

    with tab2:
        uploaded_files = st.file_uploader("광고 이미지 업로드", type=["png", "jpg", "jpeg"], accept_multiple_files=True, key=f"{prefix}_uploader_{selected_comp}")
        if uploaded_files:
            st.session_state[sess_key] = [{"id": f"upload_{i}_{len(f.read())}", "fn": f.name, "bytes": f.getvalue()} for i, f in enumerate(uploaded_files, start=1)]

    items = st.session_state[sess_key]
    if items:
        st.divider()
        st.markdown(f"**수집된 소재 ({len(items)}건) — 타겟 연령대가 다른 소재는 ❌ 삭제하세요**")
        items_to_remove = []
        for i in range(0, len(items), 6):
            row_items = items[i:i + 6]
            grid_cols = st.columns(6)
            for idx, item in enumerate(row_items):
                with grid_cols[idx]:
                    if item.get("bytes"): st.image(item["bytes"], width=110)
                    if st.button("❌", key=f"del_{item['id']}_{selected_comp}"): items_to_remove.append(item['id'])
        if items_to_remove:
            st.session_state[sess_key] = [it for it in st.session_state[sess_key] if it['id'] not in items_to_remove]
            st.rerun()
        st.info("✅ 이 소재는 세션에 저장되어 있어요. 다른 경쟁사도 이어서 수집한 뒤, **03 · 메시지 갭 분석** 탭에서 한 번에 통합 분석하시면 됩니다.")


# ------------------------------------------------------------------
# 상단 헤더 & 멀티 AI 엔진 선택 영역
# ------------------------------------------------------------------
top_col1, top_col2, top_col3 = st.columns([2.5, 1.2, 1.5])
with top_col1:
    st.markdown('<div class="appbar"><div class="appbar-emoji">🍀</div><div><div class="appbar-title">경쟁사 광고 소재 분석</div></div></div>', unsafe_allow_html=True)

with top_col2:
    ai_provider = st.selectbox("AI 엔진 선택", ["Gemini (Google)", "ChatGPT (OpenAI)", "Claude (Anthropic)"], key="selected_ai_provider")
    st.session_state["current_ai_provider"] = ai_provider

with top_col3:
    default_key = ""
    try:
        if ai_provider == "Gemini (Google)" and "GEMINI_API_KEY" in st.secrets: default_key = st.secrets["GEMINI_API_KEY"]
        elif ai_provider == "ChatGPT (OpenAI)" and "OPENAI_API_KEY" in st.secrets: default_key = st.secrets["OPENAI_API_KEY"]
        elif ai_provider == "Claude (Anthropic)" and "ANTHROPIC_API_KEY" in st.secrets: default_key = st.secrets["ANTHROPIC_API_KEY"]
    except Exception: pass

    input_api_key = st.text_input(f"{ai_provider} API Key", value=default_key, type="password", key="main_ai_api_key_input")
    st.session_state["current_api_key"] = input_api_key

st.divider()

# ------------------------------------------------------------------
# 사이드바
# ------------------------------------------------------------------
NAV_ITEMS = ["🏆 01 · 경쟁사 소재 분석", "🏠 02 · 자사 소재 분석", "🔍 03 · 메시지 갭 분석", "🎬 04 · 스토리보드 아이디어", "🗂️ 05 · 히스토리"]

# [수정] '다른 탭으로 이동' 버튼은 위젯이 이미 그려진 뒤 nav_selector 값을 직접 바꾸면
# StreamlitAPIException이 나기 때문에, 위젯을 만들기 '전'에 대기 중인 이동 요청을 먼저 반영합니다.
if "_pending_nav" in st.session_state:
    st.session_state["nav_selector"] = st.session_state.pop("_pending_nav")

with st.sidebar:
    st.markdown('<div class="sidebar-caption">SEGMENT</div>', unsafe_allow_html=True)
    segment = st.radio("사업 구분", SEGMENTS, label_visibility="collapsed", key="segment_selector", horizontal=True)
    st.divider()
    st.markdown('<div class="sidebar-caption">MENU</div>', unsafe_allow_html=True)
    st.markdown('<div class="nav-block">', unsafe_allow_html=True)
    nav_full = st.radio("메뉴", NAV_ITEMS, label_visibility="collapsed", key="nav_selector")
    st.markdown('</div>', unsafe_allow_html=True)
    nav = nav_full.split(" · ", 1)[1] if " · " in nav_full else nav_full
    nav = "0" + nav if False else nav  # (내부 로직 호환을 위해 아래에서 원래 라벨과 매칭)

# 기존 코드의 "01 · ..." 형태 라벨과 매칭되도록 변환
_NAV_MAP = {item: item.split(" ", 1)[1] for item in NAV_ITEMS}
nav = _NAV_MAP[nav_full]

if "work" not in st.session_state: st.session_state.work = {}
if segment not in st.session_state.work:
    _persisted = load_work_state(segment)
    st.session_state.work[segment] = {
        "insight": _persisted.get("insight", ""),
        "gap_analysis": _persisted.get("gap_analysis", ""),
        "ideas": _persisted.get("ideas", ""),
    }
W = st.session_state.work[segment]

def _persist_work_state():
    """인사이트 / 갭분석 / 스토리보드는 새로고침·재접속해도 남도록 즉시 저장"""
    save_work_state(segment, {
        "insight": W.get("insight", ""),
        "gap_analysis": W.get("gap_analysis", ""),
        "ideas": W.get("ideas", ""),
    })

# ------------------------------------------------------------------
# 01 · 경쟁사 소재 분석
# ------------------------------------------------------------------
if nav == "01 · 경쟁사 소재 분석":
    section_header("01", f"{segment} 경쟁사 광고 소재 분석", "경쟁사를 하나 또는 여러 개 선택하세요. 선택한 만큼만 아래에 수집창이 나타납니다.")

    competitors = load_competitors()[segment]

    # [추가] 브랜드당 최대 수집 개수 - 기존엔 12건으로 고정되어 있어서 실제 운영 소재가
    # 더 많아도 못 가져왔습니다. 이제 직접 조절할 수 있습니다.
    max_items_setting = st.slider(
        "브랜드당 최대 수집 개수", min_value=6, max_value=60, value=30, step=6,
        key=f"{segment}_max_items", help="너무 크게 잡으면 페이지 스크롤/이미지 다운로드 시간이 늘어납니다.",
    )

    add_col1, add_col2 = st.columns([4, 1.2])
    # [수정 - 3번] 멀티셀렉트는 라벨(제목 텍스트)이 있고 팝오버 버튼은 없어서 줄이 안 맞았던 문제 수정
    with add_col1:
        default_sel = st.session_state.get(f"{segment}_selected_competitors")
        if default_sel is None:
            default_sel = [competitors[0]] if competitors else []
        selected_list = st.multiselect(
            "분석할 경쟁사 선택 (여러 개 선택 가능)", competitors,
            default=[c for c in default_sel if c in competitors],
            key=f"{segment}_selected_competitors",
        )
    with add_col2:
        st.markdown('<div class="align-bottom-btn"></div>', unsafe_allow_html=True)
        with st.popover("+ 새 경쟁사 추가", use_container_width=True):
            st.markdown("##### 새 경쟁사 추가")
            new_comp = st.text_input("경쟁사명 입력", key=f"{segment}_new_comp_input")
            if st.button("추가 완료", key=f"{segment}_new_comp_btn", type="primary", use_container_width=True):
                if new_comp.strip():
                    add_competitor(segment, new_comp.strip())
                    st.success(f"'{new_comp.strip()}' 추가 완료!")
                    time.sleep(0.5)
                    st.rerun()

    default_list = DEFAULT_COMPETITORS.get(segment, [])
    custom_added = [c for c in competitors if c not in default_list]
    if custom_added:
        st.markdown("<div style='font-size: 12px; color: #555866; margin-bottom: 8px;'>📌 내가 추가한 경쟁사 (❌를 누르면 바로 삭제됩니다)</div>", unsafe_allow_html=True)
        del_cols = st.columns(min(len(custom_added), 4))
        for c_idx, c_name in enumerate(custom_added):
            with del_cols[c_idx % len(del_cols)]:
                if st.button(f"{c_name} ❌", key=f"quick_del_{segment}_{c_name}", use_container_width=True):
                    remove_competitor(segment, c_name)
                    st.success(f"'{c_name}' 삭제 완료!")
                    time.sleep(0.3)
                    st.rerun()

    if not competitors:
        st.info("등록된 경쟁사가 없습니다. 위 '+ 새 경쟁사 추가'로 먼저 추가해주세요.")
    elif not selected_list:
        st.info("경쟁사를 1개 이상 선택해주세요.")
    else:
        # 2개 이상 선택했을 때만 '한 번에 자동 수집' 버튼 노출 (1개면 아래 개별 수집으로 충분)
        if len(selected_list) > 1:
            if st.button(f"🚀 선택한 {len(selected_list)}개 경쟁사 전체 자동 수집 실행", type="primary", key=f"{segment}_bulk_crawl_btn"):
                with st.status(f"{len(selected_list)}개 경쟁사 순차 수집 중...", expanded=True) as bulk_status:
                    for c in selected_list:
                        url = META_URL_MAP.get(c, "")
                        if not url:
                            st.warning(f"'{c}'는 등록된 메타 URL이 없어 건너뜁니다. (아래 개별 칸의 '예비 업로드' 탭으로 직접 추가해주세요)")
                            continue
                        def _bulk_cb(msg, _c=c):
                            bulk_status.update(label=f"[{_c}] {msg}")
                        ads, debug_info, err = scrape_meta_ads_with_playwright(
                            url, max_items=max_items_setting, selectors=load_selectors(), status_callback=_bulk_cb
                        )
                        if err:
                            st.error(f"❌ '{c}' 수집 실패: {err}")
                        elif not ads:
                            st.info(f"⚠️ '{c}' 활성 광고를 찾지 못했습니다.")
                        else:
                            st.session_state[f"comp_items_{c}"] = ads
                            vcount = debug_info.get("video_count", 0)
                            msg = f"✅ '{c}' 이미지 소재 {len(ads)}건 수집 완료"
                            if vcount: msg += f" (동영상 {vcount}건 제외)"
                            st.success(msg)
                    bulk_status.update(label="전체 일괄 수집 완료! 아래에서 각 브랜드별 이미지를 확인하세요.", state="complete")
            st.divider()

        # [수정] 선택된 경쟁사 수만큼 수집창을 그대로 반복 표시 → 일괄수집 직후에도
        # 모든 브랜드의 이미지가 바로 아래에 보여서 '수집됐는지 안됐는지' 헷갈리지 않습니다.
        for c in selected_list:
            st.markdown(f"#### 🏷️ {c}")
            auto_url = META_URL_MAP.get(c, "")
            render_material_section("comp", c, auto_url, max_items=max_items_setting)
            st.divider()

    # 지금까지 세션에 수집해둔 경쟁사들 현황 요약
    _comp_materials, _ = gather_collected_materials(segment)
    if _comp_materials:
        summary = " · ".join([f"{name} {len(imgs)}건" for name, imgs in _comp_materials.items()])
        st.caption(f"📦 지금까지 수집된 경쟁사 소재 (전체): {summary}")

# ------------------------------------------------------------------
# 02 · 자사 소재 분석
# ------------------------------------------------------------------
elif nav == "02 · 자사 소재 분석":
    section_header("02", f"{segment} 자사 광고 소재 분석", "여기서도 소재만 모아두세요. 실제 AI 분석은 03번 탭에서 경쟁사와 함께 한 번에 진행됩니다.")
    render_material_section("own", f"자사({segment})", OWN_META_URL_MAP.get(segment, ""))

# ------------------------------------------------------------------
# 03 · 메시지 갭 분석
# ------------------------------------------------------------------
elif nav == "03 · 메시지 갭 분석":
    section_header("03", f"{segment} 메시지 갭 분석 & 위닝 포인트",
                   "경쟁사 여러 브랜드를 하나로, 자사 소재도 하나로 인식시켜 AI 호출 단 1번으로 통합 분석합니다.")

    comp_materials, own_bytes = gather_collected_materials(segment)
    comp_count = len(comp_materials)
    comp_img_total = sum(len(v) for v in comp_materials.values())

    status_cols = st.columns(2)
    with status_cols[0]:
        if comp_materials:
            st.success(f"✅ 경쟁사 소재 준비됨: {comp_count}개 브랜드 · 이미지 총 {comp_img_total}건")
        else:
            st.warning("⚠️ 01 탭에서 경쟁사 소재를 먼저 수집해주세요.")
    with status_cols[1]:
        if own_bytes:
            st.success(f"✅ 자사 소재 준비됨: 이미지 {len(own_bytes)}건")
        else:
            st.warning("⚠️ 02 탭에서 자사 소재를 먼저 수집해주세요. (없어도 경쟁사 트렌드만 볼 수 있어요)")

    st.divider()

    if not comp_materials and not own_bytes:
        st.info("먼저 01·02 탭에서 소재를 수집해주세요.")
    elif comp_materials and own_bytes:
        if st.button("🔍 경쟁사 + 자사 통합 분석 실행 (AI 요청 1회)", type="primary", key="combined_gap_btn"):
            if not st.session_state.get("current_api_key"):
                st.error("상단에서 API Key를 입력해 주세요.")
            else:
                comp_flat = [b for imgs in comp_materials.values() for b in imgs]
                with st.status("경쟁사+자사 통합 분석 진행 중...", expanded=True) as gap_status:
                    def _update_gap_status(msg):
                        gap_status.update(label=msg, state="running")
                    try:
                        resp_text = run_combined_gap_analysis(
                            st.session_state["current_ai_provider"],
                            st.session_state["current_api_key"],
                            comp_flat, own_bytes,
                            list(comp_materials.keys()), f"자사({segment})",
                            status_callback=_update_gap_status,
                        )
                        gap_status.update(label="통합 분석 완료!", state="complete")
                        W["gap_analysis"] = resp_text
                        W["insight"] = ""  # 이제 인사이트도 통합 리포트 안에 포함되므로 별도 필드는 비워둠
                        _persist_work_state()
                    except QuotaExceededError as e:
                        gap_status.update(label="무료 API 사용량 한도 초과", state="error")
                        st.error(f"🚦 {e}")
                        st.info(
                            "💡 (1) 1분 정도 기다렸다가 다시 눌러보세요. "
                            "(2) 계속 반복되면 Google AI Studio에서 결제(유료 티어) 전환을 고려해보세요. "
                            "(3) 상단 'AI 엔진 선택'에서 ChatGPT/Claude로 잠시 바꿔서 시도해보세요."
                        )
                    except Exception as e:
                        gap_status.update(label="분석 중 오류 발생", state="error")
                        st.error(f"분석 중 오류 발생: {e}")
    elif comp_materials and not own_bytes:
        st.markdown("**자사 소재가 아직 없어요 — 우선 경쟁사 트렌드만 뽑아볼 수 있어요.**")
        if st.button("경쟁사 트렌드만 분석 (AI 요청 1회)", key="comp_only_btn"):
            if not st.session_state.get("current_api_key"):
                st.error("상단에서 API Key를 입력해 주세요.")
            else:
                comp_flat = [b for imgs in comp_materials.values() for b in imgs]
                with st.status("경쟁사 소재 트렌드 분석 진행 중...", expanded=True) as trend_status:
                    def _update_trend_status(msg):
                        trend_status.update(label=msg, state="running")
                    try:
                        resp_text = run_competitor_trend_only(
                            st.session_state["current_ai_provider"], st.session_state["current_api_key"],
                            comp_flat, list(comp_materials.keys()), status_callback=_update_trend_status,
                        )
                        trend_status.update(label="분석 완료!", state="complete")
                        W["insight"] = resp_text
                        _persist_work_state()
                    except QuotaExceededError as e:
                        trend_status.update(label="무료 API 사용량 한도 초과", state="error")
                        st.error(f"🚦 {e}")
                    except Exception as e:
                        trend_status.update(label="분석 중 오류 발생", state="error")
                        st.error(f"분석 중 오류 발생: {e}")

    if W.get("gap_analysis"):
        st.divider()
        st.markdown("### 📊 경쟁사 + 자사 통합 분석 리포트")
        st.markdown(W["gap_analysis"])
        st.divider()
        if st.button("📝 이 결과로 신규 소재 아이디어 만들기 →", type="primary", key="jump_to_04"):
            st.session_state["_pending_nav"] = "🎬 04 · 스토리보드 아이디어"
            st.rerun()
    elif W.get("insight"):
        st.divider()
        st.markdown("### 📊 경쟁사 트렌드 리포트")
        st.markdown(W["insight"])

# ------------------------------------------------------------------
# 04 · 스토리보드 아이디어
# ------------------------------------------------------------------
elif nav == "04 · 스토리보드 아이디어":
    section_header("04", f"{segment} 맞춤형 스토리보드 아이디어", "브랜드/제품명만 입력해도 바로 생성돼요. 나머지 정보는 선택이에요 (채우면 더 정교해집니다).")

    all_brands = load_all_brands()
    saved_brand = all_brands.get(segment, {})
    for fkey, default in [
        ("brand_name", ""), ("brand_product", ""), ("brand_usp", ""),
        ("target_audience", ""), ("brand_design_memory", ""),
    ]:
        skey = f"{segment}_{fkey}"
        if skey not in st.session_state:
            st.session_state[skey] = saved_brand.get(fkey, default)

    # [수정] 필수 입력은 브랜드/제품명 1개만 - 나머지는 선택(펼쳐보기)로 축소해서 진입 장벽을 낮췄습니다.
    brand_name = st.text_input("브랜드/제품명 *", key=f"{segment}_brand_name", placeholder="예: 브랜드명 입력")

    with st.expander("선택 입력 (없어도 진행돼요 - 채우면 더 정교한 아이디어가 나옵니다)", expanded=False):
        col_a, col_b = st.columns(2)
        with col_a:
            brand_usp = st.text_area("핵심 셀링포인트 (USP)", key=f"{segment}_brand_usp", placeholder="예: 우리 제품만의 차별점")
            target_audience = st.text_input("타겟 고객층", key=f"{segment}_target_audience", placeholder="예: 학부모 / 자녀 연령대 등")
        with col_b:
            brand_product = st.text_area("제품/서비스 설명", key=f"{segment}_brand_product", placeholder="비워두면 03번 자사 소재 분석 내용으로 자동 추정합니다")
            brand_design_memory = st.text_area(
                "디자인 톤앤매너 메모리", key=f"{segment}_brand_design_memory",
                placeholder="예: 밝고 친근한 톤, 과장된 비교 광고는 지양함",
            )
        if st.button("이 정보 저장해두기", key=f"{segment}_brand_save"):
            save_brand(segment, {
                "brand_name": brand_name, "brand_product": brand_product, "brand_usp": brand_usp,
                "target_audience": target_audience, "brand_design_memory": brand_design_memory,
            })
            st.success("저장되었습니다. 다음에 이 부문(세그먼트)을 열면 자동으로 채워져요.")

    st.divider()

    comp_materials, own_bytes = gather_collected_materials(segment)

    if not brand_name.strip():
        st.warning("브랜드/제품명만 입력하면 바로 생성할 수 있어요. ⬆️ 위 칸을 채워주세요.")
    elif not comp_materials:
        st.info("먼저 01 탭에서 경쟁사 소재를 수집해 주세요.")
    else:
        if not W.get("gap_analysis") and not W.get("insight"):
            st.info("💡 03 탭에서 통합 분석을 먼저 실행하면 훨씬 더 정교한 스토리보드가 나와요. (지금도 생성은 가능합니다)")
        if st.button("위닝 스토리보드 아이디어 생성", type="primary"):
            if not st.session_state.get("current_api_key"):
                st.error("상단에서 API Key를 입력해 주세요.")
            else:
                gap_context = W.get("gap_analysis") or W.get("insight") or "없음 (아직 03 탭에서 통합 분석을 실행하지 않음)"
                # [수정] 제품 설명을 안 채웠으면 세그먼트명으로 기본값 대체
                effective_product = brand_product.strip() if brand_product.strip() else f"{segment} 대상 학습 서비스/제품"

                STORYBOARD_PROMPT = """당신은 크리에이티브 디렉터입니다. 아래 경쟁사 분석 결과, 메시지 갭 분석, 우리 브랜드 정보,
그리고 자사 디자인 메모리를 반영하여 차별화된 **광고 크리에이티브 스토리보드 3개**를 제안해주세요.
브랜드 정보 중 비어있는 항목은 메시지 갭 분석 내용을 참고해서 합리적으로 추정해 진행해주세요.

[자사 브랜드 정보]
- 브랜드/제품명: {brand_name}
- 제품 설명: {brand_product}
- 핵심 USP: {brand_usp}
- 타겟 고객: {target_audience}
- 자사 디자인 가이드: {design_memory}

[메시지 갭 분석]
{gap_context}

각 아이디어는 아래 구조의 스토리보드 형식으로 작성해주세요:
### [아이디어 N] 한줄 컨셉 타이틀
- **타겟구간 / 매체 소구 포인트**: 
- **훅킹 카피 (오프닝 3초)**: 
- **비주얼 구성안 (연출 기획)**: 
- **본문 설득 및 USP 소구 방식**: 
- **CTA (행동 유도 문구)**: 
- **차별화 포인트**: 
"""
                with st.status("스토리보드 기획안 작성 중...", expanded=True) as story_status:
                    def _update_story_status(msg):
                        story_status.update(label=msg, state="running")
                    try:
                        resp_text = run_unified_ai_prompt(
                            st.session_state["current_ai_provider"],
                            st.session_state["current_api_key"],
                            STORYBOARD_PROMPT.format(
                                brand_name=brand_name, brand_product=effective_product,
                                brand_usp=brand_usp or "미입력 (갭 분석 참고해서 추정)",
                                target_audience=target_audience or "미입력 (세그먼트 특성 기준으로 추정)",
                                design_memory=brand_design_memory or "기본 톤앤매너",
                                gap_context=gap_context
                            ),
                            status_callback=_update_story_status,
                        )
                        story_status.update(label="스토리보드 완성!", state="complete")
                        W["ideas"] = resp_text
                        _persist_work_state()

                        save_history_entry({
                            "timestamp": now_kst().strftime("%Y-%m-%d %H:%M"),
                            "segment": segment,
                            "brand_name": brand_name,
                            "brand_product": effective_product,
                            "target_audience": target_audience,
                            "material_count": len(comp_materials) + (1 if own_bytes else 0),
                            "competitor_names": list(comp_materials.keys()),  # [추가] 어떤 경쟁사를 분석했는지 기록
                            "insight": W.get("insight", ""),
                            "gap_analysis": W.get("gap_analysis", ""),
                            "ideas": W["ideas"],
                        })
                    except QuotaExceededError as e:
                        story_status.update(label="무료 API 사용량 한도 초과", state="error")
                        st.error(f"🚦 {e}")
                    except Exception as e:
                        story_status.update(label="오류 발생", state="error")
                        st.error(f"오류 발생: {e}")

        if W["ideas"]:
            st.markdown(W["ideas"])
            st.divider()
            st.markdown("**📥 01~04 전체 내용 문서로 저장**")
            _brand_info = {
                "brand_name": brand_name, "brand_product": brand_product, "brand_usp": brand_usp,
                "target_audience": target_audience, "brand_design_memory": brand_design_memory,
            }
            render_export_buttons(segment, _brand_info, W.get("insight", ""), W.get("gap_analysis", ""), W["ideas"], key_prefix="tab04")
            st.download_button(
                "📝 아이디어만 마크다운(.md)으로 저장", data=W["ideas"],
                file_name=f"{segment}_ad_winning_storyboards.md", mime="text/markdown",
                key="tab04_md_dl",
            )

# ------------------------------------------------------------------
# 05 · 히스토리
# ------------------------------------------------------------------
elif nav == "05 · 히스토리":
    section_header("05", "히스토리", "완료한 아이디어 추출 결과가 부문 구분과 함께 자동으로 쌓입니다.")

    history = load_history()
    show_all = st.checkbox("모든 부문 보기 (선택 해제 시 현재 부문만)", value=False)
    filtered = history if show_all else [h for h in history if h.get("segment", "") == segment]

    if not filtered:
        st.info("아직 완료된 결과가 없어요. 04 탭에서 아이디어를 생성하면 여기에 자동으로 기록됩니다.")
    else:
        col_h1, col_h2 = st.columns([3, 1])
        with col_h1: st.caption(f"총 {len(filtered)}건의 기록")
        with col_h2:
            if st.button("전체 기록 삭제"):
                save_json(HISTORY_FILE, [])
                if os.path.exists(HISTORY_FILE): os.remove(HISTORY_FILE)
                st.rerun()

        for i, entry in enumerate(filtered):
            title = f"[{entry.get('segment', '-')}] {entry.get('brand_name', '(브랜드명 없음)')} · {entry.get('timestamp', '')}"
            with st.expander(title):
                comp_names = entry.get("competitor_names", [])
                own_included = entry.get("material_count", 0) > len(comp_names)
                comp_names_str = ", ".join(comp_names) if comp_names else "기록 없음"
                st.markdown(
                    f'<div class="history-meta">분석 경쟁사 {len(comp_names)}개 '
                    f'({comp_names_str}){" · 자사 소재 포함" if own_included else ""} '
                    f'· 타겟: {entry.get("target_audience", "-")}</div>',
                    unsafe_allow_html=True,
                )
                st.divider()
                if entry.get("insight"):
                    st.markdown("**위닝 포인트 요약**")
                    st.markdown(entry["insight"])
                    st.divider()
                if entry.get("gap_analysis"):
                    st.markdown("**메시지 갭 분석**")
                    st.markdown(entry["gap_analysis"])
                    st.divider()
                st.markdown("**스토리보드 아이디어**")
                st.markdown(entry.get("ideas", ""))
                st.divider()
                st.caption("📥 이 기록 문서로 저장")
                render_export_buttons(
                    entry.get("segment", segment),
                    {"brand_name": entry.get("brand_name", ""), "brand_product": entry.get("brand_product", ""),
                     "target_audience": entry.get("target_audience", "")},
                    entry.get("insight", ""), entry.get("gap_analysis", ""), entry.get("ideas", ""),
                    key_prefix=f"history_{i}",
                )
                st.download_button(
                    "📝 아이디어만 마크다운(.md)으로 저장", data=entry.get("ideas", ""),
                    file_name=f"ad_ideas_{entry.get('timestamp', '').replace(':', '').replace(' ', '_')}.md",
                    mime="text/markdown", key=f"history_dl_{i}",
                )
